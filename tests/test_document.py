"""Регрессии: пустые разделы, координаты и содержимое Word-документов."""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

from tz_reviewer.document import load_text, numbered_document, split_sections


def test_empty_parent_and_final_step_are_preserved():
    text = "# Спецификация\n## Алгоритм\n### Чтение\nЧитаем данные.\n### Запись\n"
    sections = split_sections(text)
    assert [s.title for s in sections] == ["Спецификация", "Алгоритм", "Чтение", "Запись"]
    assert [s.body for s in sections] == ["", "", "Читаем данные.", ""]
    assert [s.parent_start_line for s in sections] == [None, 1, 2, 2]
    assert [s.level for s in sections] == [1, 2, 3, 3]
    prompt_text = numbered_document(sections)
    assert "[Раздел 4: Запись]" in prompt_text
    assert prompt_text.count("Читаем данные.") == 1


def test_body_coordinates_preserve_original_indentation_and_internal_blank_lines():
    text = "Введение\n\n## Источник\n\n  первая строка  \n\n\tвторая строка\n\n## Выход"
    sections = split_sections(text)
    source = sections[1]
    assert (source.start_line, source.end_line, source.body_start_line) == (3, 8, 5)
    assert source.body == "  первая строка  \n\n\tвторая строка"
    lines = text.splitlines()
    assert "\n".join(lines[source.body_start_line - 1:source.body_start_line + 2]) == source.body
    assert sections[0].body == "Введение"
    assert sections[-1].body_start_line is None
    assert sections[-1].end_line == 9


def test_numbered_hierarchy_returns_to_parent_for_siblings():
    sections = split_sections("1. Обработка\n1.1 Чтение\nТекст\n1.2 Запись\n2. Контроли")
    assert [s.number for s in sections] == ["1", "1.1", "1.2", "2"]
    assert [s.parent_start_line for s in sections] == [None, 1, 1, None]
    assert [s.level for s in sections] == [1, 2, 2, 1]


@pytest.mark.parametrize("fence", ["```", "~~~~"])
def test_headings_inside_fenced_examples_are_not_sections(fence):
    code = f"{fence}text\n# example heading\nSELECT\n1. Example\n{fence}"
    sections = split_sections(f"## Пример данных\n{code}\n## DDL")
    assert [s.title for s in sections] == ["Пример данных", "DDL"]
    assert sections[0].body == code


def test_shorter_fence_does_not_end_code_block():
    text = "## Пример\n````\n```\n# still code\n````\n## DDL"
    assert [s.title for s in split_sections(text)] == ["Пример", "DDL"]


@pytest.mark.parametrize("text", ["", "\n \n", "Обычный текст без заголовков."])
def test_document_without_headings_has_single_section(text):
    sections = split_sections(text)
    assert len(sections) == 1
    assert sections[0].body == text.strip()
    assert sections[0].level == 0


def test_uppercase_empty_heading_is_preserved():
    sections = split_sections("ИСТОЧНИКИ\n\nПРИЕМНИКИ\nНе применимо")
    assert len(sections) == 2
    assert sections[0].body == ""
    assert sections[1].body == "Не применимо"


def test_real_mart_empty_write_step_reaches_llm_input():
    path = Path(__file__).resolve().parents[1] / "examples" / "case_3_device_agg_vitrina.md"
    sections = split_sections(load_text(path))
    write_step = next(s for s in sections if s.title == "Шаг 5. Запись в CDM")
    assert write_step.body == ""
    parent = next(s for s in sections if s.start_line == write_step.parent_start_line)
    assert parent.title == "Алгоритм расчёта"
    assert "Шаг 5. Запись в CDM]" in numbered_document(sections)


def _add_hyperlink(paragraph, label, url=None, anchor=None):
    hyperlink = OxmlElement("w:hyperlink")
    if url:
        rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), rel_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_docx_keeps_paragraph_table_order_links_and_empty_headings(tmp_path):
    document = Document()
    document.add_heading("Источники", level=1)
    paragraph = document.add_paragraph("Каталог: ")
    _add_hyperlink(paragraph, "Карточка", "https://catalog.example/objects/arbitrary-stream")
    paragraph.add_run(" — описание.")
    table = document.add_table(rows=2, cols=3)
    for cell, value in zip(table.rows[0].cells, ["Источник", "Ссылка", "Описание"]):
        cell.text = value
    table.cell(1, 0).text = "source_xyz"
    _add_hyperlink(table.cell(1, 1).paragraphs[0], "Объект", "https://catalog.example/objects/other")
    table.cell(1, 2).text = "A | B\nСледующая строка"
    document.add_paragraph("После таблицы.")
    document.add_heading("Запись", level=2)
    path = tmp_path / "spec.docx"
    document.save(path)

    text = load_text(path)
    assert "Каталог: [Карточка](<https://catalog.example/objects/arbitrary-stream>) — описание." in text
    assert "[Объект](<https://catalog.example/objects/other>)" in text
    assert "A \\| B<br>Следующая строка" in text
    assert text.index("Каталог:") < text.index("| Источник") < text.index("После таблицы.")
    sections = split_sections(text)
    assert [s.title for s in sections] == ["Источники", "Запись"]
    assert sections[-1].body == ""
    assert sections[-1].parent_start_line == sections[0].start_line


def test_docx_nested_table_and_internal_links_are_not_lost(tmp_path):
    document = Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Вложенные сведения"
    nested = cell.add_table(rows=2, cols=1)
    nested.cell(0, 0).text = "Объект"
    nested.cell(1, 0).text = "nested_source"
    _add_hyperlink(document.add_paragraph(), "Перейти [DDL]", anchor="ddl")
    path = tmp_path / "nested.docx"
    document.save(path)
    text = load_text(path)
    assert "nested_source" in text
    assert "[Перейти \\[DDL\\]](<#ddl>)" in text


def test_docx_merged_cells_and_blank_values_keep_table_columns(tmp_path):
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Источники"
    table.cell(1, 0).text = "source_without_link"
    path = tmp_path / "merged.docx"
    document.save(path)
    text = load_text(path)
    assert "| Источники | Источники |" in text
    assert "| source_without_link |  |" in text
